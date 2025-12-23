import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import io
import logging

# --- Logging for Production Debugging ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("KeyMatcher")

def extract_matches_from_pdf(pdf_file, keyword):
    """Search PDF and return paragraphs containing the keyword."""
    results = []
    try:
        pdf_file.seek(0)
        pdf_bytes = pdf_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        for page_num, page in enumerate(doc):
            blocks = page.get_text("blocks")
            for b in blocks:
                text = b[4]
                if keyword.lower() in text.lower():
                    clean_text = " ".join(text.split())
                    results.append(f"(Page {page_num + 1}): {clean_text}")
        doc.close()
        return results
    except Exception as e:
        logger.error(f"PDF Analysis Error: {e}")
        return None

def merge_to_template(word_file, matches, keyword):
    """Appends occurrences to Word doc with safety fallback for styles."""
    try:
        word_file.seek(0)
        doc = Document(io.BytesIO(word_file.read()))
        
        doc.add_page_break()
        doc.add_heading(f"Extracted Occurrences: '{keyword}'", level=1)
        
        if not matches:
            doc.add_paragraph(f"No occurrences of '{keyword}' found in the PDF.")
        else:
            # Check if 'List Bullet' style exists in this specific document
            has_bullet_style = 'List Bullet' in doc.styles
            
            for match in matches:
                if has_bullet_style:
                    doc.add_paragraph(match, style='List Bullet')
                else:
                    # Fallback: Manual bullet point if style is missing
                    doc.add_paragraph(f"• {match}")

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream
    except Exception as e:
        # This will now print the exact error to your terminal for debugging
        logger.error(f"Word Update Error: {e}")
        return None

# --- UI Layer ---
st.set_page_config(page_title="PDF Key Matcher", layout="wide")
st.title("Enterprise PDF to Word Key Matcher")

col1, col2 = st.columns(2)
with col1:
    pdf_upload = st.file_uploader("Source PDF", type=["pdf"])
with col2:
    word_upload = st.file_uploader("Destination Word Doc (Template)", type=["docx"])

key = st.text_input("Enter the 'key' word to extract:")

if st.button("Process and Merge", use_container_width=True):
    if not (pdf_upload and word_upload and key):
        st.error("Please provide both files and a keyword.")
    else:
        with st.spinner("Analyzing PDF and updating document..."):
            found_matches = extract_matches_from_pdf(pdf_upload, key)
            
            if found_matches is None:
                st.error("Could not read the PDF.")
            else:
                final_doc = merge_to_template(word_upload, found_matches, key)
                
                if final_doc:
                    st.success(f"Successfully processed! Found {len(found_matches)} matches.")
                    st.download_button(
                        label="📥 Download Updated Word Document",
                        data=final_doc,
                        file_name=f"Updated_{word_upload.name}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("Error generating the Word file. Check terminal for details.")
